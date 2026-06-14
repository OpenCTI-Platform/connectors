from io import BytesIO
from unittest.mock import MagicMock

from docx import Document
from reportimporter.constants import (
    MIME_DOCX,
    OBSERVABLE_CLASS,
    OBSERVABLE_DETECTION_CUSTOM_REGEX,
    RESULT_FORMAT_CATEGORY,
    RESULT_FORMAT_MATCH,
    RESULT_FORMAT_TYPE,
)
from reportimporter.models import Observable
from reportimporter.report_parser import ReportParser

IPV4_STIX_TARGET = "IPv4-Addr.value"
IPV4_REGEX = r"\b(?:\d{1,3}\.){3}\d{1,3}\b"


def _build_parser() -> ReportParser:
    ipv4_observable = Observable(
        name="IPv4 test",
        detection_option=OBSERVABLE_DETECTION_CUSTOM_REGEX,
        regex_patterns=IPV4_REGEX,
        stix_target=IPV4_STIX_TARGET,
    )
    return ReportParser(
        helper=MagicMock(),
        entity_list=[],
        observable_list=[ipv4_observable],
    )


def _docx_bytes(paragraphs=(), table_rows=()) -> BytesIO:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=0, cols=cols)
        for row in table_rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def test_docx_mime_type_is_registered():
    parser = _build_parser()
    assert parser.supported_file_types[MIME_DOCX] == parser._parse_docx


def test_parse_docx_extracts_observables_from_paragraphs_and_tables():
    parser = _build_parser()
    file_data = _docx_bytes(
        paragraphs=["Reach the C2 at 8.8.8.8 for details.", "", "Nothing here."],
        table_rows=[["Indicator", "Value"], ["c2", "1.1.1.1"]],
    )

    result = parser._parse_docx(file_data)

    assert "8.8.8.8" in result
    assert "1.1.1.1" in result
    assert result["8.8.8.8"][RESULT_FORMAT_TYPE] == OBSERVABLE_CLASS
    assert result["8.8.8.8"][RESULT_FORMAT_CATEGORY] == IPV4_STIX_TARGET
    assert result["1.1.1.1"][RESULT_FORMAT_MATCH] == "1.1.1.1"


def test_parse_docx_without_iocs_returns_empty():
    parser = _build_parser()
    file_data = _docx_bytes(paragraphs=["Just some prose without any indicator."])

    assert parser._parse_docx(file_data) == {}


def test_parse_docx_with_corrupt_data_is_handled_gracefully():
    parser = _build_parser()

    # Not a valid .docx (zip) payload — the parser must swallow the error.
    assert parser._parse_docx(BytesIO(b"not a docx file")) == {}
