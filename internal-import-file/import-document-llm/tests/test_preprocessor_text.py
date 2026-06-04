"""Unit tests for non-OCR preprocessor paths and helpers."""

import os
import sys
from io import BytesIO

import docx

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../src")))

from reportimporter import preprocessor
from reportimporter.preprocessor import (
    FilePreprocessor,
    _coerce_bool,
    _process_docx,
    _try_decode,
    _unfold_cert_blocks,
)


class TestCoerceBool:
    def test_none_uses_default(self):
        assert _coerce_bool(None, default=True) is True
        assert _coerce_bool(None, default=False) is False

    def test_passthrough_bool(self):
        assert _coerce_bool(True, default=False) is True
        assert _coerce_bool(False, default=True) is False

    def test_string_values(self):
        assert _coerce_bool("true", default=False) is True
        assert _coerce_bool("1", default=False) is True
        assert _coerce_bool("on", default=False) is True
        assert _coerce_bool("false", default=True) is False
        assert _coerce_bool("0", default=True) is False


class TestTryDecode:
    def test_utf8(self):
        assert _try_decode("héllo".encode("utf-8")) == "héllo"

    def test_latin1_fallback(self):
        # 0xff is invalid utf-8 but decodes under latin-1
        assert _try_decode(b"\xff") is not None


def test_unfold_cert_blocks_joins_issuer():
    text = "Issuer: Example CA\n    CN=Root"
    assert _unfold_cert_blocks(text) == "Issuer: Example CA CN=Root"


class TestPreprocessFile:
    def test_markdown(self):
        out = FilePreprocessor.preprocess_file(b"# Title", "text/markdown", "x.md")
        assert "# Title" in out

    def test_html(self):
        out = FilePreprocessor.preprocess_file(b"<p>Hello</p>", "text/html", "x.html")
        assert "Hello" in out

    def test_csv(self):
        out = FilePreprocessor.preprocess_file(b"h1,h2\n1,2", "text/csv", "x.csv")
        assert "| 1 | 2 |" in out

    def test_plain_text(self):
        out = FilePreprocessor.preprocess_file(b"hello world", "text/plain", "x.txt")
        assert "hello world" in out

    def test_unknown_mime_falls_back_to_decode(self):
        out = FilePreprocessor.preprocess_file(b"raw data", "application/zip", "x.zip")
        assert "raw data" in out

    def test_returns_none_on_undecodable_pdf(self):
        # Not a valid PDF -> _process_pdf returns None -> preprocess_file None
        out = FilePreprocessor.preprocess_file(b"not a pdf", "application/pdf", "x.pdf")
        assert out is None


class TestProcessDocx:
    def _make_docx(self) -> bytes:
        document = docx.Document()
        document.add_heading("Main Title", level=1)
        document.add_paragraph("A paragraph of body text.")
        document.add_paragraph("Bullet item", style="List Bullet")
        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()

    def test_process_docx_markdown(self):
        out = _process_docx(self._make_docx())
        assert "# Main Title" in out
        assert "A paragraph of body text." in out
        assert "- Bullet item" in out

    def test_process_docx_invalid_returns_none(self):
        assert _process_docx(b"not a docx") is None

    def test_preprocess_file_docx_path(self):
        out = FilePreprocessor.preprocess_file(
            self._make_docx(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "report.docx",
        )
        assert "# Main Title" in out


def test_set_helper_injects_global():
    class Dummy:
        connector_logger = object()

    sentinel = Dummy()
    preprocessor.set_helper(sentinel)
    assert preprocessor._helper is sentinel
    # restore null helper to avoid leaking into other tests
    preprocessor.set_helper(preprocessor._NullHelper())
